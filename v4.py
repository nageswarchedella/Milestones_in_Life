from tkinter import *
import pathlib
from tkinter import messagebox,filedialog
from datetime import date,timedelta
import docx

def generateData(name,download_path,birthdate):
	try:
		doc=docx.Document()
		[dd,mm,yyyy]=list(map(int,birthdate.split('/')))
		d1 = date(yyyy,mm,dd)
		d2 = date.today()
		daysold=str((d2 - d1).days)
		dweekday=d1.weekday()
		weekdays=['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
		doc.add_paragraph('Welcome to Milestones in Life! \nThanks '+name +' for choosing us!\nWishing you many more Happy moments return on your Milestones.')
		doc.add_paragraph("Date of birth: " + str(dd)+"/"+str(mm)+"/"+str(yyyy))
		doc.add_paragraph("Number of days old on "+str(date.today())+" is "+str(daysold))
		doc.add_paragraph("Day of the birth: "+ str(weekdays[dweekday]))
		dya=list(range(1000,51000,1000))
		for i in dya:
			year=timedelta(days=i)
			dweekday=(d1+year).weekday()
			doc.add_paragraph("Milestone of " + str(i)+ ' days on '+str(d1+year)+' on '+str(weekdays[dweekday]))
		doc.add_paragraph("No of days old on your 100 birthdays")
		for i in range(1,101):
			d3 = date(yyyy+i,mm,dd)
			idaysold=str((d3 - d1).days)
			if str(i)[-1]=='1':
				suffix='st'
			elif str(i)[-1]=='2':
				suffix='nd'
			elif str(i)[-1]=='3':
				suffix='rd'
			else:
				suffix='th'
			doc.add_paragraph("No of Days old on your "+str(i)+suffix+' birthday is ' + str(idaysold))
		doc.save(download_path+"/"+name+".docx")
		return True
	except:
		return False

def Widgets():
	link_label=Label(root,text="Enter Name",bg='#E8D579',width=30)
	link_label.grid(row=1,column=0,pady=10, padx=10)

	linkText = Entry(root,width=55,textvariable=username)
	linkText.grid(row=1,column=1,padx=10,pady=10,columnspan=5)

	date_label=Label(root,text="Enter date of birth: dd/mm/yyyy",bg='#E8D579',width=30)
	date_label.grid(row=2,column=0,pady=10, padx=10)

	dateText = Entry(root,width=55,textvariable=dateofbirth)
	dateText.grid(row=2,column=1,padx=10,pady=10,columnspan=5)

	destination_label = Label(root,text="Destination :",bg='#E8D579',width=30)
	destination_label.grid(row=3,column=0,pady=10,padx=10)

	destinationText = Entry(root,width=40,textvariable=download_Path)
	destinationText.grid(row=3,column=1,pady=10, padx=10)

	browse_B = Button(root, text="Browse", command=Browse,width=15,bg='#05E8E0')
	browse_B.grid(row=3,column=2,padx=1,pady=1)

	Download_B = Button(root, text="Download", command=Download,width=20,bg='#05E8E0')
	Download_B.grid(row=4,column=1,pady=0,padx=0)

def Browse():
	download_Directory = filedialog.askdirectory(initialdir=pathlib.Path.cwd())
	download_Path.set(download_Directory)

def Download():
	userName = username.get()
	download_Folder = download_Path.get()
	birthdate = dateofbirth.get()
	if generateData(userName,download_Folder,birthdate):
		messagebox.showinfo("","Sucessfully downloaded")
	else:
		messagebox.showwarning("Data is invalid","Date is incorrect or any field is invalid. Please try again!")


root = Tk()
root.geometry("640x200")
root.resizable(0,0)
root.title("MileStones in life")
root.config(background='#000000')
username = StringVar()
download_Path = StringVar()
dateofbirth = StringVar()

Widgets()
root.mainloop()
