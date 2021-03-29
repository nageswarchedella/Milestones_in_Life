from datetime import date,timedelta
import docx
doc=docx.Document()
datetrue=True
print("Welcome to Milestones in Life")
name=input("Enter your name:")
while datetrue:
	try:
		[dd,mm,yyyy]=list(map(int,input("Enter you Birthday as day-month-year: ").split('-')))
		d1 = date(yyyy,mm,dd)
		d2 = date.today()
		daysold=str((d2 - d1).days)
		dweekday=d1.weekday()
		weekdays=['Monday','Tuesday','Wednesday','Thursday','Friday','Saturday','Sunday']
		dya=list(range(1000,51000,1000))
		doc.add_paragraph('Welcome to Milestones in Life! \nThanks '+name +' for choosing us!\nWishing you many more Happy moments return on your Milestones.')
		doc.add_paragraph("Number of days old on "+str(date.today())+" is "+str(daysold))
		doc.add_paragraph("Day of the birth: "+ str(weekdays[dweekday]))
		print("Generating your Milestones...",end='')
		for i in dya:
			year=timedelta(days=i)
			dweekday=(d1+year).weekday()
			doc.add_paragraph("Milestone of " + str(i)+ ' days on '+str(d1+year)+' on '+str(weekdays[dweekday]))
		doc.add_paragraph("No of days old on your birthdays")
		for i in range(1,100):
			d3 = date(yyyy+i,mm,dd)
			daysold=str((d3 - d1).days)
			if str(i)[-1]=='1':
				suffix='st'
			elif str(i)[-1]=='2':
				suffix='nd'
			elif str(i)[-1]=='3':
				suffix='rd'
			else:
				suffix='th'
			doc.add_paragraph("No of Days old on your "+str(i)+suffix+' birthday is ' + str(daysold))
		doc.save("F:/practice problems/Python/milestones in life/"+name+".docx")
		print("Completed")
		print("Thanks for using Milestones in Life")
		datetrue=False
	except ValueError:
		print('Date is invalid.. Try again')
